# SPDX-License-Identifier: MIT
# Copyright (c) 2025 Scott Joiner

import os
import argparse
import pickle
import tempfile
from docx import Document
from docxcompose.composer import Composer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request

SCOPES = ['https://www.googleapis.com/auth/drive.file']


def collect_lesson_files(folder_path, sort_by='name'):
    lesson_files = []
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(".docx"):
                lesson_files.append(os.path.join(root, file))

    if sort_by == 'ctime':
        lesson_files.sort(key=lambda f: os.path.getctime(f))
    else:
        lesson_files.sort(key=lambda f: os.path.basename(f).lower())

    return lesson_files


def add_table_of_contents(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def prepare_docs_with_breaks(docx_files):
    temp_paths = []
    for i, path in enumerate(docx_files):
        doc = Document(path)
        if i != len(docx_files) - 1:
            doc.add_page_break()
        temp_fd, temp_path = tempfile.mkstemp(suffix=".docx", prefix=f"lesson_{i}_")
        os.close(temp_fd)
        doc.save(temp_path)
        temp_paths.append(temp_path)
    return temp_paths


def merge_with_images(docx_files, output_filename):
    if not docx_files:
        print("❌ No .docx files found to merge.")
        return None

    print(f"Merging {len(docx_files)} files into {output_filename}...")
    temp_paths = prepare_docs_with_breaks(docx_files)

    master = Document(temp_paths[0])
    composer = Composer(master)

    for path in temp_paths[1:]:
        composer.append(Document(path))

    add_table_of_contents(master)
    composer.save(output_filename)
    print(f"✅ Merged lessons into: {output_filename}")

    # Cleanup temp files
    for temp_file in temp_paths:
        try:
            os.remove(temp_file)
        except Exception as e:
            print(f"⚠️ Could not delete temp file {temp_file}: {e}")

    return output_filename


def upload_to_google_drive(filepath):
    creds = None
    if os.path.exists('token.pickle'):
        with open('token.pickle', 'rb') as token:
            creds = pickle.load(token)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not os.path.exists('credentials.json'):
                print("❌ Missing 'credentials.json'. Download it from Google Cloud Console.")
                return
            flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)

        with open('token.pickle', 'wb') as token:
            pickle.dump(creds, token)

    service = build('drive', 'v3', credentials=creds)

    file_metadata = {
        'name': os.path.splitext(os.path.basename(filepath))[0],
        'mimeType': 'application/vnd.google-apps.document'
    }
    media = MediaFileUpload(filepath, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    uploaded = service.files().create(body=file_metadata, media_body=media, fields='id,webViewLink').execute()
    print(f"\n📤 Uploaded to Google Docs: {uploaded.get('webViewLink')}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Merge and upload DOCX lessons to Google Drive")
    parser.add_argument("folder_path", help="Path to folder with lesson .docx files (recursive)")
    parser.add_argument("--merged_name", default=None, help="Filename for merged output (default: foldername.docx)")
    parser.add_argument("--sort", choices=["name", "ctime"], default="name",
                        help="How to sort lessons: 'name' or 'ctime' (default: name)")
    args = parser.parse_args()

    merged_name = args.merged_name or os.path.basename(os.path.normpath(args.folder_path)) + ".docx"
    lesson_paths = collect_lesson_files(args.folder_path, sort_by=args.sort)
    merged_file = merge_with_images(lesson_paths, merged_name)

    if merged_file:
        upload_to_google_drive(merged_file)
