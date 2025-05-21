# SPDX-License-Identifier: MIT
# Copyright (c) 2025 Scott Joiner

# Needs to run first
from harmony_tools import config

import os
import shutil
import click
import requests
import re
import base64
import tempfile
import cairosvg
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


# --- Helper functions ---
def safe_filename(name):
    return "".join(c for c in name if c.isalnum() or c in " -_").rstrip()


def svg_to_png_bytes(svg_html):
    """
    Convert an SVG string to PNG bytes in memory.
    Attempts to auto-correct missing namespace issues.
    """
    try:
        # Fix missing xmlns if necessary
        if "<svg" in svg_html and "xmlns=" not in svg_html:
            svg_html = svg_html.replace(
                "<svg",
                '<svg xmlns="http://www.w3.org/2000/svg" xmlns:xlink="http://www.w3.org/1999/xlink"',
                1,
            )

        png_data = cairosvg.svg2png(bytestring=svg_html.encode("utf-8"))
        return png_data

    except Exception as e:
        print(f"Failed to convert SVG to PNG: {e}")
        return None


def download_image(url, save_folder):
    try:
        response = requests.get(url, stream=True, timeout=10)
        response.raise_for_status()
        parsed_url = urlparse(url)
        img_name = os.path.basename(parsed_url.path)
        img_name = safe_filename(img_name)
        img_path = os.path.join(save_folder, img_name)
        with open(img_path, "wb") as f:
            f.write(response.content)
        return img_path
    except Exception as e:
        print(f"Failed to download image {url}: {e}")
        return None


def add_hyperlink(paragraph, text, url):
    """
    A function that places a hyperlink within a paragraph object.
    """
    # Create the w:hyperlink tag and add needed values
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a w:r element
    new_run = OxmlElement("w:r")

    # Create a w:rPr element
    rPr = OxmlElement("w:rPr")

    # Add color
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    # Add underline
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    # Add rPr to run
    new_run.append(rPr)

    # Create a w:t element and add the text
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    # DONE! Do not add (url) visibly after


def process_element(elem, doc):
    """
    Main Processing Loop — handles block-level structures.
    """
    if elem.name in ["h1", "h2", "h3"]:
        handle_heading(elem, doc)

    elif elem.name == "p":
        handle_paragraph(elem, doc)

    elif elem.name == "ul":
        handle_unordered_list(elem, doc)

    elif elem.name == "ol":
        handle_ordered_list(elem, doc)

    elif elem.name == "img":
        # Standalone block-level image
        para = doc.add_paragraph()
        handle_image(elem, para, doc)

    elif elem.name == "div":
        elem_classes = elem.get("class", [])

        if "lecture-attachment-type-pdf_embed" in elem_classes:
            handle_pdf_embed(elem, doc)

        elif "lecture-attachment-type-audio" in elem_classes:
            audio_name = "[Audio]"
            name_span = elem.find("span", class_="audioloader__name")
            if name_span and name_span.string:
                audio_name = f"[{name_span.string.strip()}]"
            doc.add_paragraph(audio_name)

        elif "lecture-attachment-type-video" in elem_classes:
            doc.add_paragraph("[Video Here]")

        else:
            # Default: recurse into div's children
            for child in elem.find_all(recursive=False):
                process_element(child, doc)

    else:
        # Unknown elements are recursively processed
        for child in elem.find_all(recursive=False):
            process_element(child, doc)


def handle_heading(elem, doc):
    # Determine heading level (limit to Heading 1–3 for Word styles)
    heading_level = min(int(elem.name[1]), 3)

    # Create the heading paragraph
    para = doc.add_paragraph(style=f"Heading {heading_level}")

    # Process inline contents (text, links, images inside heading)
    process_inline_contents(elem, para, doc)

    # Optional: add a little spacing after heading
    para.paragraph_format.space_after = Pt(6)


def handle_paragraph(elem, doc, style="Normal"):
    para = doc.add_paragraph(style=style)
    process_inline_contents(elem, para, doc)

    if para:
        if style != "Normal":
            para.paragraph_format.left_indent = Inches(0.5)

        para.paragraph_format.space_after = Pt(10 if style == "Normal" else 4)


def handle_unordered_list(elem, doc):
    for li in elem.find_all("li", recursive=False):
        handle_paragraph(li, doc, style="List Bullet")


def handle_ordered_list(elem, doc):
    for li in elem.find_all("li", recursive=False):
        handle_paragraph(li, doc, style="List Number")


def insert_png_into_paragraph(png_bytes, para, width_inches=0.3):
    if not png_bytes:
        print("No PNG bytes — skipping insertion.")
        return

    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
        tmp.write(png_bytes)
        tmp_path = tmp.name

    print(f"Saved temporary PNG at {tmp_path}")  # <<== NEW

    run = para.add_run()
    run.add_picture(tmp_path, width=Inches(width_inches))

    # Comment out delete for now
    os.remove(tmp_path)


def process_inline_contents(elem, para, doc, bold=False, italic=False):
    for child in elem.contents:
        if isinstance(child, NavigableString):
            text = child.strip()
            if text and para:
                run = para.add_run(text)
                # run.font.size = Pt(11)
                run.bold = bold
                run.italic = italic

        elif child.name == "a" and child.has_attr("href"):
            link_text = child.get_text(strip=True) or child["href"]
            href = child["href"]
            if para:
                para.add_run(" ")
                add_hyperlink(para, link_text, href)
                para.add_run(" ")

        elif child.name == "strong":
            if para:
                para.add_run(" ")
                process_inline_contents(child, para, doc, bold=True, italic=italic)
                para.add_run(" ")

        elif child.name == "em":
            if para:
                para.add_run(" ")
                process_inline_contents(child, para, doc, bold=bold, italic=True)
                para.add_run(" ")

        elif child.name in ["span", "br"]:
            process_inline_contents(child, para, doc, bold=bold, italic=italic)

        elif child.name == "img":
            if para:
                handle_image(child, para, doc)

        elif child.name == "svg":
            # if para:
            #    svg_html = str(child)
            #    png_bytes = svg_to_png_bytes(svg_html)
            #    if png_bytes:
            #        width_inches, height_inches = extract_svg_dimensions(svg_html)
            #        insert_png_into_paragraph(png_bytes, para, width_inches, height_inches)
            continue

        elif child.name in ["math", "canvas"]:
            continue  # Ignore weird inline things

        else:
            # BLOCK ELEMENT FOUND — flush para and delegate to block handler
            para = None
            process_element(child, doc)


def parse_style_for_maxwidth(style_string):
    """
    Extract max-width in pixels from a style attribute string.
    """
    if not style_string:
        return None

    match = re.search(r"max-width:\s*(\d+)px", style_string)
    if match:
        try:
            return int(match.group(1))
        except:  # noqa: E722
            return None

    return None


def parse_style_for_dimensions(style_string):
    """
    Parses the style attribute and returns width_px, max_width_px if found.
    """
    width = None
    max_width = None

    if not style_string:
        return width, max_width

    styles = style_string.split(";")
    for item in styles:
        key_value = item.split(":")
        if len(key_value) != 2:
            continue
        key = key_value[0].strip().lower()
        value = key_value[1].strip()

        if key == "width" and value.endswith("px"):
            try:
                width = int(value.replace("px", "").strip())
            except Exception:
                pass
        elif key == "max-width" and value.endswith("px"):
            try:
                max_width = int(value.replace("px", "").strip())
            except Exception:
                pass

    return width, max_width


def extract_image_dimensions(img_elem):
    """
    Extract width and height attributes from an <img> tag.
    Return (width_inches, height_inches).
    """
    width = None
    height = None

    width_attr = img_elem.get("width")
    height_attr = img_elem.get("height")
    style_attr = img_elem.get("style")

    width_style, max_width_style = parse_style_for_dimensions(style_attr)

    if width_attr:
        width = int(width_attr) / 96
    elif width_style:
        width = int(width_style) / 96
    elif max_width_style:
        width = int(max_width_style) / 96

    if height_attr:
        height = int(height_attr) / 96

    return width, height


def handle_image(elem, para, doc):
    img_src = elem.get("src")
    if not img_src:
        return

    try:
        if img_src.startswith("data:image"):
            header, base64_data = img_src.split(",", 1)
            image_data = base64.b64decode(base64_data)

            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                tmp.write(image_data)
                tmp_path = tmp.name

            temp_file_created = True

        else:
            img_url = urljoin(f"file://{doc.input_path}", img_src)
            tmp_path = download_image(img_url, doc.images_folder)
            temp_file_created = False

        if tmp_path:
            width_inches, height_inches = extract_image_dimensions(elem)

            # Only center if para is otherwise empty (block images)
            if not para.runs:  # Meaning no prior text in the paragraph
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = para.add_run()

            if width_inches or height_inches:
                run.add_picture(
                    tmp_path,
                    width=Inches(width_inches) if width_inches else None,
                    height=Inches(height_inches) if height_inches else None,
                )
            else:
                # No size info, use 100% usable width
                section = para.part.document.sections[0]
                page_width = section.page_width
                left_margin = section.left_margin
                right_margin = section.right_margin
                usable_width = page_width - left_margin - right_margin
                run.add_picture(tmp_path, width=usable_width)

            if temp_file_created:
                os.remove(tmp_path)

    except Exception as e:
        print(f"Failed to insert image {img_src}: {e}")


def handle_pdf_embed(elem, doc):
    # Try to find the download block
    label_div = elem.find("div", class_="label")
    if label_div:
        pdf_title = label_div.get_text(strip=True)
    else:
        pdf_title = "PDF Attachment"

    download_link = None
    download_a = elem.find("a", href=True)
    if download_a:
        download_link = download_a["href"]

    # Insert into Word
    para = doc.add_paragraph(style="Normal")
    run = para.add_run(f"📎 Attached Document: {pdf_title}")
    run.bold = True

    if download_link:
        para = doc.add_paragraph(style="Normal")
        para.add_run("Download here: ")
        add_hyperlink(para, download_link, download_link)


def process_file(filename):
    input_path = os.path.join(config.INPUT_FOLDER, filename)

    with open(input_path, "r", encoding="utf-8") as file:
        soup = BeautifulSoup(file, "html5lib")

    body = soup.body
    if not body:
        print(f"Warning: No <body> tag found in {filename}. Skipping.")
        return

    page_title = (
        soup.title.string.strip()
        if soup.title and soup.title.string
        else filename.replace(".html", "")
    )
    lesson_folder_name = safe_filename(page_title)
    lesson_folder = os.path.join(config.OUTPUT_FOLDER, lesson_folder_name)
    os.makedirs(lesson_folder, exist_ok=True)
    images_folder = os.path.join(lesson_folder, "images")
    os.makedirs(images_folder, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Helvetica"

    lesson_body = soup.find("div", class_="course-mainbar lecture-content")
    if not lesson_body:
        print(f"Warning: No lesson body found in {filename}. Skipping.")
        return

    print(
        f"Found {len(lesson_body.find_all('div', class_='lecture-attachment'))} lecture-attachment blocks."
    )
    block_counter = 0
    for block in lesson_body.find_all(recursive=False):
        block_counter += 1
        print(f"--- Processing content block {block_counter}---")

        if block.name in ["script", "meta", "style"]:
            continue

        first_child = block.find(recursive=False)
        if (
            first_child
            and first_child.has_attr("class")
            and "comments" in first_child["class"]
        ):
            print("Skipping comment-only block.")
            continue

        process_element(block, doc)

    output_path = os.path.join(lesson_folder, f"{lesson_folder_name}.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")

    processed_path = os.path.join(config.PROCESSED_FOLDER, filename)
    shutil.move(input_path, processed_path)
    print(f"Moved: {filename} -> Processed folder")


@click.command(help="Convert HTML lessons to DOCX")
@click.option(
    "--nomedia", is_flag=True, default=False, help="Skip downloading and embed images"
)
@click.option("--font", default=None, help="Override default font Helvetica")
@click.option("--workdir", default=None, help="Override default working directory")
def main(nomedia, font, workdir):
    config.init(workdir)

    for filename in os.listdir(config.INPUT_FOLDER):
        if filename.lower().endswith(".html"):
            process_file(filename)
    print("\nAll lessons processed!")


if __name__ == "__main__":
    print("🔔 RUNNING HTML --> DOCX SCRIPT - VERSION B!")
    main()
